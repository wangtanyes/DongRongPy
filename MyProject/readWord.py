#读取docx中的文本代码示例
import docx
#获取文档对象
file=docx.Document("D:\\文件\\20180104-中银国际-泰胜风能（300129.SZ）：增资昌力科技，军工转型第一步.docx")
print("段落数:"+str(len(file.paragraphs)))#段落数为13，每个回车隔离一段

#输出每一段的内容
for para in file.paragraphs:
    print("*****", para.text)

#输出段落编号及段落内容
for i in range(len(file.paragraphs)):
    print("第"+str(i)+"段的内容是："+file.paragraphs[i].text)
